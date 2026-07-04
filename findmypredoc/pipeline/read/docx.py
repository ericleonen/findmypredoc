import io

import docx
import requests


def read(url: str) -> str:
    """
    Downloads a DOCX file from a URL and extracts its text.
    """

    response = requests.get(url, timeout=30)
    response.raise_for_status()

    return read_bytes(response.content)


def read_bytes(data: bytes) -> str:
    """
    Extracts text from raw DOCX bytes, including paragraphs and table cells.
    """

    document = docx.Document(io.BytesIO(data))

    parts = [paragraph.text for paragraph in document.paragraphs]

    for table in document.tables:
        for row in table.rows:
            parts.append(" | ".join(cell.text for cell in row.cells))

    return "\n".join(parts)
